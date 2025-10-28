import streamlit as st
import re
import fitz
import pandas as pd
from docx import Document
from typing import List, Dict, Any
from openai import OpenAI
import json
import os
import pdfplumber
import numpy as np
import time
from datetime import datetime

# Page config
st.set_page_config(page_title="Clinical Document Processor", page_icon="🏥", layout="wide")

# API Key setup
api_key = st.secrets["api_key"]
client = OpenAI(api_key=api_key)

# ==================== DASHBOARD METRICS ====================
class ProcessingMetrics:
    def __init__(self):
        self.start_time = None
        self.protocol_start = None
        self.protocol_end = None
        self.crf_start = None
        self.crf_end = None
        self.chunks_created = 0
        self.chunks_processed = 0
        self.protocol_tables_found = 0
        self.protocol_tables_processed = 0
        self.protocol_pages_scanned = 0
        self.crf_items_extracted = 0
        
    def start_processing(self):
        self.start_time = time.time()
        
    def get_elapsed_time(self):
        if self.start_time:
            return time.time() - self.start_time
        return 0
    
    def get_protocol_time(self):
        if self.protocol_start and self.protocol_end:
            return self.protocol_end - self.protocol_start
        return 0
    
    def get_crf_time(self):
        if self.crf_start and self.crf_end:
            return self.crf_end - self.crf_start
        return 0

def format_time(seconds):
    """Format seconds to readable time"""
    if seconds < 60:
        return f"{seconds:.1f}s"
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes}m {secs}s"

def create_dashboard(metrics):
    """Create live metrics dashboard"""
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("⏱️ Total Time", format_time(metrics.get_elapsed_time()))
    with col2:
        st.metric("📄 Protocol Pages", metrics.protocol_pages_scanned)
    with col3:
        st.metric("🧩 CRF Chunks", f"{metrics.chunks_processed}/{metrics.chunks_created}")
    with col4:
        st.metric("📊 Tables Processed", metrics.protocol_tables_processed)
    
    # Progress bars
    if metrics.chunks_created > 0:
        chunk_progress = metrics.chunks_processed / metrics.chunks_created
        st.progress(chunk_progress, text=f"CRF Processing: {metrics.chunks_processed}/{metrics.chunks_created} chunks")
    
    # Detailed stats in expander
    with st.expander("📈 Detailed Statistics"):
        stats_col1, stats_col2 = st.columns(2)
        with stats_col1:
            st.write("**Protocol Processing**")
            st.write(f"- Time: {format_time(metrics.get_protocol_time())}")
            st.write(f"- Tables Found: {metrics.protocol_tables_found}")
            st.write(f"- Tables Processed: {metrics.protocol_tables_processed}")
        with stats_col2:
            st.write("**CRF Processing**")
            st.write(f"- Time: {format_time(metrics.get_crf_time())}")
            st.write(f"- Chunks Created: {metrics.chunks_created}")
            st.write(f"- Items Extracted: {metrics.crf_items_extracted}")

# ==================== CORE FUNCTIONS ====================

def map_data_manually(source_df, header_row=4, start_row=5):
    """Maps data from source to template"""
    MANUAL_COLUMN_MAP = {
        'form_label':'Form Label',
        'item_group':'Item Group (if only one on form, recommend same as Form Label)',
        'item_group_repeating':'Item group Repeating',
        'item_order':'Item Order',
        'item_label':'Item Label',
        'item_name':'Item Name (provided by SDTM Programmer, if SDTM linked item)',
        'data_type':'Data type',
        'codelist':'Codelist - Choice Labels (If many, can use Codelist Tab)',
        'codelist_name':'Codelist Name (provided by SDTM programmer)',
        'required':'Required'
    }
    
    try:
        print(f"[MAP] Mapping {len(source_df)} rows to template")
        req_headers = list(MANUAL_COLUMN_MAP.values())
        template_headers = [
            'New or Copied from Study','Form Label','Form Name(provided by SDTM Programmer, if SDTM linked form)',
            'Item Group (if only one on form, recommend same as Form Label)','Item group Repeating',
            'Repeat Maximum, if known, else default =50','Display format of repeating item group (Grid, read only, form)',
            'Default Data in repeating item group','Item Order','Item Label',
            'Item Name (provided by SDTM Programmer, if SDTM linked item)','Progressively displayed?',
            'Controlling Item (item name if known, else item label)','Controlling Item Value','','Data type',
            'If text or number, Field Length','If number, Precision (decimal places)',
            'Codelist - Choice Labels (If many, can use Codelist Tab)','Codelist Name (provided by SDTM programmer)',
            'Choice Code (provided by SDTM programmer)','Codelist: Control Type','If Number, Range: Min Value - Max Value',
            "If Date, Query Future Date",'Required','If Required, Open Query when Intentionally Left Blank (form, item)','Notes'
        ]
        
        template_file = pd.DataFrame(columns=template_headers)
        
        # Check columns
        source_columns_in_map = list(MANUAL_COLUMN_MAP.keys())
        missing_cols = set(source_columns_in_map) - set(source_df.columns)
        if missing_cols:
            print(f"[MAP] Warning: Missing columns {missing_cols}")
        
        # Map columns
        df_selected = source_df[source_columns_in_map]
        df_mapped = df_selected.rename(columns=MANUAL_COLUMN_MAP)
        template_file[req_headers] = df_mapped
        
        print(f"[MAP] Successfully mapped {len(template_file)} rows")
        return template_file
        
    except Exception as e:
        print(f"[MAP] Error: {e}")
        return pd.DataFrame()

def table_ai(combined_data, system_prompt):
    """Clean table data using AI"""
    combined_data2 = [{'data': combined_data.to_json(orient='records')}]
    user_prompt_pr = f"""INPUT JSON: {combined_data2}
    Clean and return the structured JSON"""
    
    messages_new = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': user_prompt_pr}
    ]
    
    try:
        print(f"[AI] Cleaning table with {len(combined_data)} rows")
        response = client.chat.completions.create(
            model="o4-mini",
            messages=messages_new,
            response_format={"type": "json_object"},
        )
        
        cleaned_data_json = json.loads(response.choices[0].message.content)
        
        if 'data' in cleaned_data_json and cleaned_data_json['data']:
            result = pd.DataFrame(cleaned_data_json['data'])
            print(f"[AI] Cleaned table: {len(result)} rows")
            return result
        elif cleaned_data_json:
            return pd.DataFrame(cleaned_data_json)
        else:
            print("[AI] Warning: Empty response from API")
            return pd.DataFrame()
            
    except Exception as e:
        print(f"[AI] Error: {e}")
        return pd.DataFrame()

def combine_rows(df3):
    """Combine rows with forward and backward fill"""
    print(f"[COMBINE] Processing {len(df3)} rows")
    fd = pd.DataFrame()
    df3[0] = df3[0].fillna(method='ffill')
    groups = df3[0].unique().tolist()
    
    for i in range(len(groups)):
        try:
            group_df = df3[df3[0] == groups[i]]
            df1 = group_df.ffill()
            df2 = group_df.bfill()
            df = pd.concat([df1, df2])
            
            result = {}
            for col in df.columns:
                values = df[col].dropna()
                unique_values = values.unique()
                
                if len(unique_values) == 0:
                    result[col] = np.nan
                elif len(unique_values) == 1:
                    result[col] = unique_values[0]
                else:
                    result[col] = ' '.join(str(v) for v in unique_values)
            
            fd = pd.concat([fd, pd.DataFrame([result])], ignore_index=True)
            
        except Exception as e:
            print(f"[COMBINE] Error processing group {groups[i]}: {e}")
    
    result = fd.drop_duplicates().fillna('').reset_index(drop=True)
    print(f"[COMBINE] Combined to {len(result)} rows")
    return result

# ==================== PROTOCOL POST-PROCESSING ====================

def parse_visit_window(window_str):
    """Parse visit window string to get numeric value"""
    try:
        if pd.isna(window_str) or window_str == '':
            return 0
        window_str = str(window_str).strip()
        if window_str.startswith('±'):
            return int(window_str.replace('±', ''))
        elif '/' in window_str:  # Handle "0/+5" case
            return window_str
        else:
            return 0
    except Exception as e:
        print(f"[PARSE_WINDOW] Error parsing window '{window_str}': {e}")
        return 0

def map_protocol_to_schedule_grid(protocol_df):
    """
    Map protocol visit details to Schedule Grid format with error handling

    Parameters:
    protocol_df: DataFrame with columns including 'Visit short name', 'Study week', 'Visit window (days)'

    Returns:
    schedule_grid_df: DataFrame in Schedule Grid format, or empty DataFrame on error
    """
    try:
        print("[SCHEDULE_GRID] Starting protocol to schedule grid mapping")

        # Validate input
        if protocol_df is None or protocol_df.empty:
            print("[SCHEDULE_GRID] Error: Empty protocol DataFrame")
            return pd.DataFrame()

        # Check if required rows exist
        # required_procedures = ['Visit short name', 'Study week', 'Visit window (days)']
        # if 'Procedure' not in protocol_df.columns:
        #     print("[SCHEDULE_GRID] Error: 'Procedure' column not found")
        #     return pd.DataFrame()

        # available_procedures = protocol_df['Procedure'].tolist()
        # missing_procedures = [p for p in required_procedures if p not in available_procedures]
        # if missing_procedures:
        #     print(f"[SCHEDULE_GRID] Error: Missing required procedures: {missing_procedures}")
        #     return pd.DataFrame()
        # Check if required rows exist with flexible pattern matching
        required_patterns = {
            'visit_short_name': ['visit short name', 'visit name', 'short name', 'visit id'],
            'study_week': ['study week', 'week', 'study day', 'timepoint'],
            'visit_window': ['visit window', 'window', 'window days', 'visit window (days)']
        }
        
        if 'Procedure' not in protocol_df.columns:
            print("[SCHEDULE_GRID] Error: 'Procedure' column not found")
            return pd.DataFrame()
        
        # Find matches for each required pattern
        found_rows = {}
        missing_patterns = []
        
        for key, patterns in required_patterns.items():
            found = False
            for pattern in patterns:
                for idx, value in enumerate(protocol_df['Procedure']):
                    if pd.notna(value) and pattern.lower() in str(value).lower():
                        found_rows[key] = idx
                        found = True
                        break
                if found:
                    break
            if not found:
                missing_patterns.append(key)
        
        if missing_patterns:
            print(f"[SCHEDULE_GRID] Error: Missing required patterns: {missing_patterns}")
            return pd.DataFrame()

# Access matched rows using: found_rows['visit_short_name'], found_rows['study_week'], etc.

        # Extract relevant rows from protocol dataframe
        visit_names = protocol_df[protocol_df['Procedure'] == 'Visit short name'].iloc[0, 2:].values
        study_weeks = protocol_df[protocol_df['Procedure'] == 'Study week'].iloc[0, 2:].values
        visit_windows = protocol_df[protocol_df['Procedure'] == 'Visit window (days)'].iloc[0, 2:].values

        # Initialize schedule grid data
        schedule_data = {
            'Form Label': [],
            'RTSM': [],
            'Screening': [],
            'Randomisation': []
        }

        # Process each visit to determine which category it falls into
        visits_list = []
        for i, (visit_name, week, window) in enumerate(zip(visit_names, study_weeks, visit_windows)):
            if pd.notna(visit_name):
                visit_name_clean = str(visit_name).strip()
                try:
                    week_num = float(week) if pd.notna(week) else 0
                except:
                    week_num = 0

                window_val = parse_visit_window(window)

                visits_list.append({
                    'visit_name': visit_name_clean,
                    'week': week_num,
                    'window': window_val,
                    'index': i
                })

        # Validate minimum visits
        total_visits = len(visits_list)
        if total_visits < 3:
            print(f"[SCHEDULE_GRID] Error: Protocol must have at least 3 visits (found {total_visits})")
            return pd.DataFrame()

        print(f"[SCHEDULE_GRID] Processing {total_visits} visits")

        # DYNAMIC CATEGORIZATION based on position
        screening_visit = visits_list[0]  # First visit
        randomisation_visit = visits_list[1]  # Second visit
        end_of_study_visit = visits_list[-1]  # Last visit
        follow_up_visit = visits_list[-2] if total_visits > 2 else None
        end_of_treatment_visit = visits_list[-3] if total_visits > 3 else None

        # Main Study: From 3rd visit to visit before EOT
        main_study_visits = []
        if total_visits > 4:
            main_study_visits = visits_list[2:-3]
        elif total_visits == 4:
            main_study_visits = []

        # Add Main Study visit columns to schedule_data
        for v in main_study_visits:
            col_name = f"Visit {v['visit_name']}"
            schedule_data[col_name] = []

        # Add remaining columns
        if end_of_treatment_visit:
            schedule_data['End of Treatment'] = []
        if follow_up_visit:
            schedule_data['Follow Up'] = []
        if end_of_study_visit:
            schedule_data['End of Study'] = []

        # Define rows for the schedule grid
        rows = [
            'Event Label:',
            'Event Name:',
            'Visit Dynamics (If Y, then Event should appear based on triggering criteria)',
            'Triggering: Event',
            'Triggering: Form',
            'Triggering: Item = Response',
            'Event Window Configuration',
            'Assign Visit Window',
            'Offset Type',
            'Offset Days',
            'Day Range - Early',
            'Day Range - Late'
        ]

        for row_name in rows:
            row_data = {'Form Label': row_name}

            # RTSM columns (appear for Screening and Randomisation)
            if row_name == 'Event Label:':
                row_data['RTSM'] = 'RTSM'
                row_data['Screening'] = 'Screening'
                row_data['Randomisation'] = 'Randomisation'
            elif row_name == 'Event Name:':
                row_data['RTSM'] = 'RTSM'
                row_data['Screening'] = 'SCRN'
                row_data['Randomisation'] = 'RAND'
            elif row_name == 'Visit Dynamics (If Y, then Event should appear based on triggering criteria)':
                row_data['RTSM'] = ''
                row_data['Screening'] = 'Y'
                row_data['Randomisation'] = 'Y'
            elif row_name == 'Triggering: Event':
                row_data['RTSM'] = ''
                row_data['Screening'] = 'SCRN'
                row_data['Randomisation'] = 'RAND'
            elif row_name == 'Triggering: Form':
                row_data['RTSM'] = ''
                row_data['Screening'] = 'ELIGIBILITY_CRITERIA'
                row_data['Randomisation'] = 'RANDOMISATION'
            elif row_name == 'Triggering: Item = Response':
                row_data['RTSM'] = ''
                row_data['Screening'] = 'DSSTDAT_ELIG'
                row_data['Randomisation'] = 'DSSTDAT_RAND'
            elif row_name == 'Assign Visit Window':
                row_data['RTSM'] = ''
                row_data['Screening'] = 'Y'
                row_data['Randomisation'] = 'Y'
            elif row_name == 'Offset Type':
                row_data['RTSM'] = ''
                row_data['Screening'] = 'Specific: SCRN'
                row_data['Randomisation'] = 'Specific: RAND'
            elif row_name == 'Offset Days':
                row_data['RTSM'] = ''
                row_data['Screening'] = '3'
                row_data['Randomisation'] = '3'
            elif row_name == 'Day Range - Early':
                row_data['RTSM'] = ''
                center_day = int(screening_visit['week'] * 7)
                if isinstance(screening_visit['window'], str) and '/' in screening_visit['window']:
                    parts = screening_visit['window'].split('/')
                    early_offset = int(parts[0]) if parts[0] else 0
                    row_data['Screening'] = str(center_day + early_offset)
                else:
                    window = screening_visit['window'] if isinstance(screening_visit['window'], int) else 14
                    row_data['Screening'] = str(center_day - window)
                row_data['Randomisation'] = '0'
            elif row_name == 'Day Range - Late':
                row_data['RTSM'] = ''
                center_day = int(screening_visit['week'] * 7)
                if isinstance(screening_visit['window'], str) and '/' in screening_visit['window']:
                    parts = screening_visit['window'].split('/')
                    late_offset = int(parts[1].replace('+', '')) if len(parts) > 1 else 14
                    row_data['Screening'] = str(center_day + late_offset)
                else:
                    window = screening_visit['window'] if isinstance(screening_visit['window'], int) else 14
                    row_data['Screening'] = str(center_day - window)
                row_data['Randomisation'] = '0'
            else:
                row_data['RTSM'] = ''
                row_data['Screening'] = ''
                row_data['Randomisation'] = ''

            # Fill Main Study visits
            for v in main_study_visits:
                col_name = f"Visit {v['visit_name']}"

                if row_name == 'Event Label:':
                    row_data[col_name] = f"Visit {v['visit_name']}"
                elif row_name == 'Event Name:':
                    row_data[col_name] = v['visit_name']
                elif row_name == 'Visit Dynamics (If Y, then Event should appear based on triggering criteria)':
                    row_data[col_name] = 'Y'
                elif row_name == 'Triggering: Event':
                    row_data[col_name] = 'AE'
                elif row_name == 'Triggering: Form':
                    row_data[col_name] = 'AE'
                elif row_name == 'Triggering: Item = Response':
                    row_data[col_name] = 'AE.AEOUT'
                elif row_name == 'Assign Visit Window':
                    row_data[col_name] = 'Y'
                elif row_name == 'Offset Type':
                    row_data[col_name] = 'Previous'
                elif row_name == 'Offset Days':
                    row_data[col_name] = '3'
                elif row_name == 'Day Range - Early':
                    center_day = int(v['week'] * 7)
                    window = v['window'] if isinstance(v['window'], int) else 3
                    row_data[col_name] = str(center_day - window)
                elif row_name == 'Day Range - Late':
                    center_day = int(v['week'] * 7)
                    window = v['window'] if isinstance(v['window'], int) else 3
                    row_data[col_name] = str(center_day + window)
                else:
                    row_data[col_name] = ''

            # Fill End of Treatment
            if end_of_treatment_visit:
                if row_name == 'Event Label:':
                    row_data['End of Treatment'] = 'End of Treatment'
                elif row_name == 'Event Name:':
                    row_data['End of Treatment'] = end_of_treatment_visit['visit_name']
                elif row_name == 'Visit Dynamics (If Y, then Event should appear based on triggering criteria)':
                    row_data['End of Treatment'] = 'Y'
                elif row_name == 'Triggering: Event':
                    row_data['End of Treatment'] = 'EOT'
                elif row_name == 'Triggering: Form':
                    row_data['End of Treatment'] = 'AE'
                elif row_name == 'Triggering: Item = Response':
                    row_data['End of Treatment'] = 'AE.AEOUT'
                elif row_name == 'Assign Visit Window':
                    row_data['End of Treatment'] = 'Y'
                elif row_name == 'Offset Type':
                    row_data['End of Treatment'] = 'Previous'
                elif row_name == 'Offset Days':
                    row_data['End of Treatment'] = '3'
                elif row_name == 'Day Range - Early':
                    center_day = int(end_of_treatment_visit['week'] * 7)
                    window = end_of_treatment_visit['window'] if isinstance(end_of_treatment_visit['window'], int) else 3
                    row_data['End of Treatment'] = str(center_day - window)
                elif row_name == 'Day Range - Late':
                    center_day = int(end_of_treatment_visit['week'] * 7)
                    window = end_of_treatment_visit['window'] if isinstance(end_of_treatment_visit['window'], int) else 3
                    row_data['End of Treatment'] = str(center_day + window)
                else:
                    row_data['End of Treatment'] = ''

            # Fill Follow Up
            if follow_up_visit:
                if row_name == 'Event Label:':
                    row_data['Follow Up'] = f'Visit {follow_up_visit["visit_name"]}'
                elif row_name == 'Event Name:':
                    row_data['Follow Up'] = follow_up_visit['visit_name']
                elif row_name == 'Visit Dynamics (If Y, then Event should appear based on triggering criteria)':
                    row_data['Follow Up'] = 'Y'
                elif row_name == 'Triggering: Event':
                    row_data['Follow Up'] = follow_up_visit['visit_name']
                elif row_name == 'Triggering: Form':
                    row_data['Follow Up'] = 'AE'
                elif row_name == 'Triggering: Item = Response':
                    row_data['Follow Up'] = 'AE.AEOUT'
                elif row_name == 'Assign Visit Window':
                    row_data['Follow Up'] = 'Y'
                elif row_name == 'Offset Type':
                    row_data['Follow Up'] = 'Previous'
                elif row_name == 'Offset Days':
                    row_data['Follow Up'] = '3'
                elif row_name == 'Day Range - Early':
                    center_day = int(follow_up_visit['week'] * 7)
                    window = follow_up_visit['window'] if isinstance(follow_up_visit['window'], int) else 3
                    row_data['Follow Up'] = str(center_day - window)
                elif row_name == 'Day Range - Late':
                    center_day = int(follow_up_visit['week'] * 7)
                    window = follow_up_visit['window'] if isinstance(follow_up_visit['window'], int) else 3
                    row_data['Follow Up'] = str(center_day + window)
                else:
                    row_data['Follow Up'] = ''

            # Fill End of Study
            if end_of_study_visit:
                if row_name == 'Event Label:':
                    row_data['End of Study'] = 'End of Study'
                elif row_name == 'Event Name:':
                    row_data['End of Study'] = end_of_study_visit['visit_name']
                elif row_name == 'Visit Dynamics (If Y, then Event should appear based on triggering criteria)':
                    row_data['End of Study'] = 'Y'
                elif row_name == 'Triggering: Event':
                    row_data['End of Study'] = 'EOS'
                elif row_name == 'Triggering: Form':
                    row_data['End of Study'] = 'AE'
                elif row_name == 'Triggering: Item = Response':
                    row_data['End of Study'] = 'AE.AEOUT'
                elif row_name == 'Assign Visit Window':
                    row_data['End of Study'] = 'Y'
                elif row_name == 'Offset Type':
                    row_data['End of Study'] = 'Previous'
                elif row_name == 'Offset Days':
                    row_data['End of Study'] = '3'
                elif row_name == 'Day Range - Early':
                    center_day = int(end_of_study_visit['week'] * 7)
                    if isinstance(end_of_study_visit['window'], str) and '/' in end_of_study_visit['window']:
                        parts = end_of_study_visit['window'].split('/')
                        early_offset = int(parts[0]) if parts[0] else 0
                        row_data['End of Study'] = str(center_day + early_offset)
                    else:
                        window = end_of_study_visit['window'] if isinstance(end_of_study_visit['window'], int) else 0
                        row_data['End of Study'] = str(center_day - window)
                elif row_name == 'Day Range - Late':
                    center_day = int(end_of_study_visit['week'] * 7)
                    if isinstance(end_of_study_visit['window'], str) and '/' in end_of_study_visit['window']:
                        parts = end_of_study_visit['window'].split('/')
                        late_offset = int(parts[1].replace('+', '')) if len(parts) > 1 else 0
                        row_data['End of Study'] = str(center_day + late_offset)
                    else:
                        window = end_of_study_visit['window'] if isinstance(end_of_study_visit['window'], int) else 0
                        row_data['End of Study'] = str(center_day + window)
                else:
                    row_data['End of Study'] = ''

            schedule_data['Form Label'].append(row_data['Form Label'])
            schedule_data['RTSM'].append(row_data.get('RTSM', ''))
            schedule_data['Screening'].append(row_data.get('Screening', ''))
            schedule_data['Randomisation'].append(row_data.get('Randomisation', ''))

            for v in main_study_visits:
                col_name = f"Visit {v['visit_name']}"
                schedule_data[col_name].append(row_data.get(col_name, ''))

            if end_of_treatment_visit:
                schedule_data['End of Treatment'].append(row_data.get('End of Treatment', ''))
            if follow_up_visit:
                schedule_data['Follow Up'].append(row_data.get('Follow Up', ''))
            if end_of_study_visit:
                schedule_data['End of Study'].append(row_data.get('End of Study', ''))

        # Create DataFrame
        schedule_grid_df = pd.DataFrame(schedule_data)
        print(f"[SCHEDULE_GRID] Successfully created schedule grid with {len(schedule_grid_df)} rows")

        return schedule_grid_df

    except Exception as e:
        print(f"[SCHEDULE_GRID] Error: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

# ==================== CRF CHUNKING ====================

class DOCXCRFChunker:
    def __init__(self, max_chunk_size: int = 15000, overlap_size: int = 500):
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size
        self.form_patterns = [
            r'sCRF\s+v\d+\.\d+',
            r'\[[\w_]+\]',
            r'V\d+(?:,\s*V\d+)*',
        ]

    def extract_and_chunk(self, docx_path: str):
        """Extract text from DOCX and create chunks"""
        print(f"[CHUNK] Starting extraction from {docx_path}")
        doc = Document(docx_path)
        elements = self._extract_structured_content(doc)
        chunks = self._create_chunks(elements)
        print(f"[CHUNK] Created {len(chunks)} initial chunks")
        return chunks

    def _extract_structured_content(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract content preserving structure"""
        elements = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = self._get_paragraph_from_element(doc, element)
                if para and para.text.strip():
                    elements.append({
                        'type': 'paragraph',
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': self._is_heading(para),
                        'is_form_title': self._is_form_title(para.text),
                        'length': len(para.text)
                    })
            elif element.tag.endswith('tbl'):
                table = self._get_table_from_element(doc, element)
                if table:
                    table_text = self._extract_table_text(table)
                    elements.append({
                        'type': 'table',
                        'text': table_text,
                        'rows': len(table.rows),
                        'cols': len(table.columns) if table.rows else 0,
                        'length': len(table_text),
                        'is_crf_table': self._is_crf_table(table_text)
                    })
        
        print(f"[CHUNK] Extracted {len(elements)} elements")
        return elements

    def _get_paragraph_from_element(self, doc, element):
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _get_table_from_element(self, doc, element):
        for table in doc.tables:
            if table._element == element:
                return table
        return None

    def _is_heading(self, para) -> bool:
        if para.style:
            style_name = para.style.name.lower()
            return 'heading' in style_name or 'title' in style_name
        return False

    def _is_form_title(self, text: str) -> bool:
        for pattern in self.form_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        form_indicators = [
            'collection of samples', 'contraception', 'c-ssrs',
            'patient health questionnaire', 'weight history',
            'hand grip test', 'mri scan consent'
        ]
        
        text_lower = text.lower()
        return any(indicator in text_lower for indicator in form_indicators)

    def _extract_table_text(self, table) -> str:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))
        return "\n".join(table_text)

    def _is_crf_table(self, table_text: str) -> bool:
        crf_indicators = [
            'yes', 'no', 'radio', 'checkbox', 'required', 'optional', '*',
            'integration', 'argus', 'cosmos', 'item label', 'data type', 'codelist'
        ]
        text_lower = table_text.lower()
        return sum(1 for indicator in crf_indicators if indicator in text_lower) >= 2

    def _create_chunks(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        chunks = []
        current_chunk = self._create_new_chunk(1)
        
        for element in elements:
            if (element['type'] == 'paragraph' and element.get('is_form_title', False)):
                if current_chunk['length'] > 0:
                    chunks.append(self._finalize_chunk(current_chunk))
                    current_chunk = self._create_new_chunk(len(chunks) + 1)
                current_chunk['form_context'] = element['text']
            
            if (current_chunk['length'] + element['length'] > self.max_chunk_size and
                current_chunk['length'] > 0):
                if element['type'] == 'table':
                    chunks.append(self._finalize_chunk(current_chunk))
                    current_chunk = self._create_new_chunk(len(chunks) + 1)
                    current_chunk['form_context'] = chunks[-1]['form_context']
            
            current_chunk['elements'].append(element)
            current_chunk['text'] += element['text'] + '\n\n'
            current_chunk['length'] += element['length']
            
            if element['type'] == 'table':
                current_chunk['has_tables'] = True
        
        if current_chunk['length'] > 0:
            chunks.append(self._finalize_chunk(current_chunk))
        
        return self._add_overlap(chunks)

    def _create_new_chunk(self, chunk_id: int) -> Dict[str, Any]:
        return {
            'chunk_id': chunk_id,
            'elements': [],
            'text': '',
            'length': 0,
            'form_context': '',
            'has_tables': False
        }

    def _finalize_chunk(self, chunk: Dict[str, Any]) -> Dict[str, Any]:
        if chunk['form_context'] and chunk['form_context'] not in chunk['text'][:200]:
            chunk['text'] = f"FORM CONTEXT: {chunk['form_context']}\n\n" + chunk['text']
        
        chunk['paragraph_count'] = sum(1 for el in chunk['elements'] if el['type'] == 'paragraph')
        chunk['table_count'] = sum(1 for el in chunk['elements'] if el['type'] == 'table')
        chunk['crf_table_count'] = sum(1 for el in chunk['elements']
                                     if el['type'] == 'table' and el.get('is_crf_table', False))
        return chunk

    def _add_overlap(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if len(chunks) <= 1:
            return chunks
        
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i-1]
            current_chunk = chunks[i]
            prev_text = prev_chunk['text']
            overlap_text = prev_text[-self.overlap_size:] if len(prev_text) > self.overlap_size else prev_text
            current_chunk['text'] = f"OVERLAP FROM PREVIOUS:\n{overlap_text}\n\nCURRENT CHUNK:\n" + current_chunk['text']
            current_chunk['has_overlap'] = True
        
        return chunks

    def get_chunk_summary(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        return {
            'total_chunks': len(chunks),
            'total_length': sum(chunk['length'] for chunk in chunks),
            'avg_chunk_size': sum(chunk['length'] for chunk in chunks) / len(chunks) if chunks else 0,
            'chunks_with_tables': sum(1 for chunk in chunks if chunk['has_tables']),
            'total_tables': sum(chunk.get('table_count', 0) for chunk in chunks),
            'total_crf_tables': sum(chunk.get('crf_table_count', 0) for chunk in chunks),
            'forms_identified': len(set(chunk['form_context'] for chunk in chunks if chunk['form_context']))
        }

def combine_chunks(chunks, max_tokens=1500):
    """Combine chunks within token limit"""
    if not chunks or len(chunks) == 0:
        return []
    
    print(f"[COMBINE_CHUNKS] Combining {len(chunks)} chunks with max_tokens={max_tokens}")
    updated_chunks = []
    current_chunk = chunks[0].copy()
    
    for i in range(1, len(chunks)):
        try:
            next_chunk = chunks[i]
            combined_length = current_chunk.get('length', 0) + next_chunk.get('length', 0) * 1.25
            
            if combined_length < max_tokens:
                current_chunk = {
                    'chunk_id': f"{current_chunk.get('chunk_id', '')}-{next_chunk.get('chunk_id', '')}",
                    'elements': current_chunk.get('elements', []) + next_chunk.get('elements', []),
                    'text': current_chunk.get('text', '') + "\n\n" + next_chunk.get('text', ''),
                    'length': current_chunk.get('length', 0) + next_chunk.get('length', 0),
                    'form_context': current_chunk.get('form_context', '') + '\n\n' + next_chunk.get('form_context', ''),
                    'has_tables': current_chunk.get('has_tables', False) or next_chunk.get('has_tables', False),
                    'paragraph_count': current_chunk.get('paragraph_count', 0) + next_chunk.get('paragraph_count', 0),
                    'table_count': current_chunk.get('table_count', 0) + next_chunk.get('table_count', 0),
                    'crf_table_count': current_chunk.get('crf_table_count', 0) + next_chunk.get('crf_table_count', 0),
                    'has_overlap': current_chunk.get('has_overlap', False) or next_chunk.get('has_overlap', False)
                }
            else:
                updated_chunks.append(current_chunk)
                current_chunk = next_chunk.copy()
                
        except Exception as e:
            print(f"[COMBINE_CHUNKS] Error at chunk {i}: {e}")
            updated_chunks.append(current_chunk)
            current_chunk = chunks[i].copy()
    
    updated_chunks.append(current_chunk)
    print(f"[COMBINE_CHUNKS] Result: {len(updated_chunks)} combined chunks")
    return updated_chunks

def process_crf_docx(docx_path: str) -> List[Dict[str, Any]]:
    """Process CRF DOCX file"""
    print(f"[CRF] Processing DOCX: {docx_path}")
    chunker = DOCXCRFChunker(max_chunk_size=15000, overlap_size=500)
    chunks = chunker.extract_and_chunk(docx_path)
    chunks = combine_chunks(chunks, 1500)
    
    summary = chunker.get_chunk_summary(chunks)
    print(f"[CRF] Created {summary['total_chunks']} chunks, {summary['forms_identified']} forms identified")
    
    return chunks

# ==================== PROTOCOL PROCESSING ====================

system_prompt_pr = """You are a clinical trial table restructuring expert. Transform the provided Schedule of Activities JSON data.

INPUT ISSUES:
- Merged visits in single cells: "V2D-2\nV2D-1 V2D1" or "V16 V17 V18"
- Merged rows: cells contain "\n" separating two row values
- Incomplete headers: row 0 has parent headers only in first column of each group
- Sometimes row 0 might not have headers (continuation from previous table) - conserve order, do not change
- Misaligned timing/window data
- CRITICAL: Row values may be shifted up or down - values that should align with visit codes might be in wrong rows
- Null values need replacement

TRANSFORMATIONS (execute in order):

0. DETECT AND CORRECT VERTICAL MISALIGNMENT - DO THIS FIRST
   - Analyze patterns in rows with similar data types (timing values, X marks, numeric values)
   - Look for patterns where values appear shifted by 1-2 rows up or down
   - Common indicators of misalignment:
     * Timing values (numbers, ±X) appearing in wrong rows relative to "Timing of Visit" label
     * X marks appearing in rows above/below procedure names
     * Empty cells where X marks should be, with X marks in adjacent rows
     * Day numbers appearing in "Visit Window" row instead of "Timing of Visit" row
   
   Detection algorithm:
   a) Identify row labels in column 0: "Visit", "Timing of Visit (Days)", "Visit Window (Days)", procedure names
   b) Check if data in columns 1+ matches expected pattern for that label
   c) If mismatch detected, check rows immediately above/below for better pattern match
   d) Shift row values up/down to correct alignment

1. UNMERGE ROWS - CRITICAL: REPLACE, DON'T DUPLICATE
   - IF a cell contains "\n", that row represents TWO rows merged
   - REMOVE the original merged row from output
   - REPLACE it with TWO separate rows
   - Split ALL cells in that row at "\n"

2. SPLIT MERGED VISITS - REPLACE, DON'T DUPLICATE
   - When ONE cell contains multiple visits (space or \n separated)
   - REMOVE that column from output
   - REPLACE with MULTIPLE columns (one per visit)
   - Distribute X marks (including Xp, Xh, Xo variants) and timing to corresponding new columns

3. RECONSTRUCT HEADERS (row 0) - ONLY IF HEADERS EXIST
   - If row 0 is blank/continuation, skip this step and preserve as-is
   - If row 0 contains parent headers spanning multiple columns:
     * Visit code format changes indicate new phase groups
     * Propagate parent header to ALL columns in its group with subscripts (_1, _2, _3)

4. PROPAGATE PHASE NAMES (column 0)
   - When column "0" is empty/null, fill with phase name from nearest row above
   - Keep phase names consistent

5. CLEAN TEXT
   - Remove ALL "\n" from final output
   - Fix broken words: "Withdraw al" → "Withdrawal"
   - Fix spellings, keep section numbers: "10.1.3", "8.1"
   - Replace None/null with ""
   - CRITICAL: Do NOT modify X mark variants - preserve exactly as-is

6. ALIGN DATA
   - After step 0 corrections, verify:
   - "Timing of Visit (Days)" row: contains day numbers
   - "Visit Window (Days)" row: contains ±N values
   - X marks in procedure rows: aligned with correct visit columns

OUTPUT:
Return ONLY this JSON (no markdown, no explanations):

{"data": [[row0_values], [row1_values], [row2_values], ...]}

Return only the JSON object."""

def extract_table_pages(pdf_file):
    """Extract Schedule of Activities pages"""
    print(f"[PROTOCOL] Extracting table pages from {pdf_file}")

    schedule_pattern = re.compile(r"schedule\s+of\s+activities", re.IGNORECASE)
    intro_pattern = re.compile(r"Introduction", re.IGNORECASE)
    index_pattern = re.compile(r"Table\s+of\s+[Cc]ontents|Contents", re.IGNORECASE)

    schedule_start_page = None
    intro_start_page = None

    pdf_document = fitz.open(pdf_file)

    print(f"[PROTOCOL] Scanning {len(pdf_document)} pages for Schedule of Activities")

    # First pass: Find schedule start page (skip only the table of contents page itself)
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text = page.get_text("text", sort=True)

        # Check if this is a Table of Contents page
        is_toc_page = index_pattern.search(text)

        # Look for Schedule of Activities (even on TOC page, but prefer actual content pages)
        if schedule_pattern.search(text):
            # If we haven't found it yet, or this is not a TOC page, use this page
            if schedule_start_page is None or not is_toc_page:
                schedule_start_page = page_num + 1
                print(f"[PROTOCOL] Schedule section found at page {schedule_start_page}")
                # If not on TOC, we can break here
                if not is_toc_page:
                    break

    # Second pass: Find introduction or end marker (only if schedule was found)
    if schedule_start_page:
        for page_num in range(schedule_start_page, len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text("text", sort=True)

            if intro_pattern.search(text):
                intro_start_page = page_num + 1
                print(f"[PROTOCOL] Introduction found at page {intro_start_page}")
                break

    if not schedule_start_page:
        st.error("❌ [PROTOCOL] Schedule of Activities section not found in the PDF")
        print("[PROTOCOL] Schedule of Activities section not found")
        pdf_document.close()
        return None

    st.write(f"📍 Found Schedule of Activities at page {schedule_start_page}")
    if intro_start_page:
        st.write(f"📍 Found Introduction at page {intro_start_page}")

    print(f"[PROTOCOL] Schedule starts at page {schedule_start_page}, Introduction at page {intro_start_page if intro_start_page else 'not found'}")
    
    end_page = intro_start_page if intro_start_page else len(pdf_document)

    table_settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "snap_tolerance": 300,
        "edge_min_length": 100,
    }

    consecutive_empty_pages = 0
    max_empty_pages = 2

    print(f"[PROTOCOL] Scanning for tables from page {schedule_start_page} to {end_page}")

    with pdfplumber.open(pdf_file) as pdf:
        for i in range(schedule_start_page - 1, len(pdf.pages)):
            page = pdf.pages[i]
            tables_on_page = page.extract_tables(table_settings=table_settings)

            # Stop if we've reached the introduction page (fix: i is 0-indexed, intro_start_page is 1-indexed)
            if intro_start_page and (i + 1) >= intro_start_page:
                print(f"[PROTOCOL] Reached introduction page, stopping at page {i + 1}")
                break
            elif tables_on_page and any(len(table) > 2 for table in tables_on_page):
                end_page = i + 1
                consecutive_empty_pages = 0
                print(f"[PROTOCOL] Found table on page {i + 1}, updating end_page to {end_page}")
            else:
                consecutive_empty_pages += 1
                print(f"[PROTOCOL] No tables on page {i + 1}, consecutive empty: {consecutive_empty_pages}")
                # If no intro marker, stop after 2 consecutive empty pages
                if consecutive_empty_pages >= max_empty_pages and not intro_start_page:
                    print(f"[PROTOCOL] Stopping after {consecutive_empty_pages} consecutive empty pages")
                    break

    st.write(f"📄 Extracting pages {schedule_start_page} to {end_page}")
    
    output_pdf = fitz.open()
    output_pdf.insert_pdf(pdf_document, from_page=schedule_start_page - 1, to_page=end_page - 1)
    # st.write('inserted pdf')
    if output_pdf.page_count > 0:
        # st.write('more than 0 pages')
        extracted_pdf_path = "Schedule_of_Activities.pdf"
        output_pdf.save(extracted_pdf_path)
        # st.write(f"[PROTOCOL] Saved {output_pdf.page_count} pages")
        pdf_document.close()
        output_pdf.close()
        return extracted_pdf_path
    else:
        output_pdf.close()
        pdf_document.close()
        st.write("[PROTOCOL] No pages extracted")
        return None

def process_protocol_pdf_pdfplumber(extracted_pdf_path, system_prompt, metrics, dashboard_placeholder):
    """Process protocol PDF with live updates"""
    print(f"[PROTOCOL] Processing PDF: {extracted_pdf_path}")
    
    if not extracted_pdf_path or not os.path.exists(extracted_pdf_path):
        print(f"[PROTOCOL] Error: File not found")
        return pd.DataFrame()
    
    table_settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "explicit_vertical_lines": [],
        "explicit_horizontal_lines": [],
        "snap_tolerance": 300,
        "snap_x_tolerance": 6,
        "snap_y_tolerance": 5.16,
        "join_tolerance": 1,
        "join_x_tolerance": 2,
        "join_y_tolerance": 3,
        "edge_min_length": 100,
        "min_words_vertical": 3,
        "min_words_horizontal": 1,
        "intersection_tolerance": 1,
        "intersection_x_tolerance": 0.4,
        "intersection_y_tolerance": 2,
        "text_tolerance": 3,
        "text_x_tolerance": 5,
        "text_y_tolerance": 3,
    }
    
    try:
        with pdfplumber.open(extracted_pdf_path) as pdf:
            # print(f"[PROTOCOL] Opened PDF with {len(pdf.pages)} pages")
            # st.write("opened file")
            df = pd.DataFrame()
            df_ai = pd.DataFrame()
            
            for i in range(len(pdf.pages)):
                page = pdf.pages[i]
                # st.write(f"[PROTOCOL] Processing page {i+1}/{len(pdf.pages)}")
                
                metrics.protocol_pages_scanned = i + 1
                dashboard_placeholder.empty()
                with dashboard_placeholder.container():
                    create_dashboard(metrics)
                
                tables_on_page = page.extract_tables(table_settings=table_settings)
                
                if tables_on_page:
                    print(f"[PROTOCOL] Found {len(tables_on_page)} tables on page {i+1}")
                    metrics.protocol_tables_found += len(tables_on_page)
                    
                    raw_data = pd.DataFrame()
                    for table_data in tables_on_page:
                        if table_data and len(table_data) >= 2:
                            raw_data = pd.concat((raw_data, pd.DataFrame(table_data)))
                    
                    if not raw_data.empty:
                        combined_data = raw_data.copy()
                        # st.write("Starting Table Cleaning")
                        nd = table_ai(combined_data, system_prompt)
                        
                        if nd is not None and not nd.empty:
                            metrics.protocol_tables_processed += 1
                            df = pd.concat((df, combined_data))
                            df_ai = pd.concat((df_ai, nd))
                            print(f"[PROTOCOL] Table cleaned successfully")
                        else:
                            metrics.protocol_tables_processed += 1
                            df = pd.concat((df, combined_data))
                            df_ai = df.copy()
                            print(f"[PROTOCOL] Table cleaned successfully")
                
                # Update dashboard
                dashboard_placeholder.empty()
                with dashboard_placeholder.container():
                    create_dashboard(metrics)
            
            all_extracted_data = df_ai
            
            if not all_extracted_data.empty:
                pr_df = pd.DataFrame(all_extracted_data)
                
                if not pr_df.empty:
                    # st.write(f"[PROTOCOL] Successfully extracted {len(pr_df)} rows")
                    pr_df.columns = pr_df.iloc[0]
                    pr_df = pr_df[1:].reset_index(drop=True)
                    pr_df = pr_df.dropna(axis=1, how='all')
                    return pr_df
            else:
                st.write("[PROTOCOL] No tables extracted")
                return pd.DataFrame()
    
    except Exception as e:
        print(f"[PROTOCOL] Error: {e}")
        return pd.DataFrame()

# ==================== CRF EXTRACTION ====================

System_prompt = """You are a CRF (Case Report Form) extraction specialist. Your task is to extract structured information from clinical research document chunks and return it as valid JSON.
REQUIRED OUTPUT FORMAT: Return ONLY a valid JSON object with the structure below. Do not include any explanatory text before or after the JSON.
{
  "forms": [
    {
      "form_label": "string",
      "form_code": "string",
      "event_codes": "string",
      "form_type": "string",
      "item_group": "string",
      "item_group_repeating": "Y/N",
      "item_order": number,
      "item_label": "string",
      "item_name": "[SDTM Programmer]",
      "data_type": "string",
      "codelist": "string",
      "codelist_name": "string",
      "control_time": "",
      "required": "Y/N"
    }
  ],
  "chunk_metadata": {
    "chunk_id": "string",
    "forms_found": number,
    "items_extracted": number
  }
}

EXTRACTION RULES:
1. Form Label: Extract the main form title (e.g., "Collection of Samples for Laboratory (Lab_1)")
2. Form Code: Extract code from Form Name field or in brackets (e.g., "[LAB_SMPL_TKN]" or "Lab_1")
3. Event Codes: Extract ALL visit codes/event names where this form appears (e.g., "V2, V12, V18" or "SCRN, V4, V5, V6"). Look for patterns like V followed by numbers, or standard visit abbreviations.
4. Form Type: "Non-repeating form" or "Repeating form"
5. Item Group: Logical sections like "Blood", "Urine", "Contraception"
6. Item Group Repeating: "Y" if section can repeat, "N" otherwise
7. Item Order: Sequential number within form (1, 2, 3...)
8. Item Label: Exact question text
9. Item Name: Always "[SDTM Programmer]"
10. Data Type: "Radio Button", "Numeric", "Text", "Date", "Checkbox"
11. Codelist: Available options (e.g., "Yes/No", "4-point Scale")
12. Codelist Name: Logical name for options
13. Required: "Y" if marked with asterisk (*), "N" otherwise

DATA TYPE MAPPING:
- Questions with radio options → "Radio Button"
- Numeric inputs (age, weight, scores) → "Numeric"
- Free text fields → "Text"
- Date fields → "Date"
- Multiple selections → "Checkbox"

Return valid JSON only. No other text."""

def user_prompt(text):
    return f"""Extract CRF information from the following document chunk and return as JSON:

CHUNK CONTENT:
{text}

Analyze this content and extract all CRF forms, item groups, and individual items following the specified JSON format. Pay special attention to:
- Form titles and codes
- Required fields marked with asterisks (*)
- Item groups and their organization
- Question text and response options
- Data types based on field characteristics

Return only valid JSON with no additional text."""

def ai_extract(chunks, system_prompt, metrics, dashboard_placeholder):
    """Extract CRF data with live updates"""
    print(f"[CRF] Starting extraction for {len(chunks)} chunks")
    
    conversation = [{"role": "system", "content": system_prompt}]
    all_forms_data = []
    
    for i, chunk in enumerate(chunks):
        print(f"[CRF] Processing chunk {i+1}/{len(chunks)}")
        
        user_msg = user_prompt(chunk['text'])
        messages = conversation + [{"role": "user", "content": user_msg}]
        
        try:
            response = client.chat.completions.create(
                model="o4-mini",
                messages=messages,
                response_format={"type": "json_object"}
            )
            answer = json.loads(response.choices[0].message.content)
            
            if 'forms' in answer:
                items_in_chunk = len(answer['forms'])
                all_forms_data.extend(answer['forms'])
                metrics.crf_items_extracted = len(all_forms_data)
                print(f"[CRF] Extracted {items_in_chunk} items from chunk {i+1}")
            
            metrics.chunks_processed = i + 1
            
            # Update dashboard
            dashboard_placeholder.empty()
            with dashboard_placeholder.container():
                create_dashboard(metrics)
            
        except Exception as e:
            print(f"[CRF] Error processing chunk {chunk['chunk_id']}: {e}")
    
    if all_forms_data:
        print(f"[CRF] Total items extracted: {len(all_forms_data)}")
        return pd.DataFrame(all_forms_data)
    else:
        print("[CRF] No forms data extracted")
        return pd.DataFrame()

# ==================== SCHEDULE GRID + CRF MERGE ====================

def create_combined_excel(schedule_grid_df, crf_df):
    """
    Create a single Excel file with two sheets:
    - 'Schedule Grid': schedule_grid_mapped data
    - 'Study Specific Forms': crf_extraction data

    Parameters:
    schedule_grid_df: DataFrame with schedule grid data
    crf_df: DataFrame with CRF extraction data (template mapped, no event_codes)

    Returns:
    bytes: Excel file content as bytes
    """
    try:
        print("[EXCEL] Creating combined Excel file")

        from io import BytesIO

        # Create a BytesIO buffer
        output = BytesIO()

        # Create Excel writer
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Write schedule grid to first sheet
            if schedule_grid_df is not None and not schedule_grid_df.empty:
                schedule_grid_df.to_excel(writer, sheet_name='Schedule Grid', index=False)
                print(f"[EXCEL] Added 'Schedule Grid' sheet with {len(schedule_grid_df)} rows")

            # Write CRF extraction to second sheet
            if crf_df is not None and not crf_df.empty:
                crf_df.to_excel(writer, sheet_name='Study Specific Forms', index=False)
                print(f"[EXCEL] Added 'Study Specific Forms' sheet with {len(crf_df)} rows")

        # Get the Excel file content
        excel_data = output.getvalue()
        print("[EXCEL] Combined Excel file created successfully")

        return excel_data

    except Exception as e:
        print(f"[EXCEL] Error creating combined Excel: {e}")
        import traceback
        traceback.print_exc()
        return None

def extract_visit_number(text):
    """
    Extract visit number from text. Handles patterns like V12, P3, Visit 5, etc.
    Returns the numeric part or None if not found.
    """
    if not text or pd.isna(text):
        return None

    text = str(text).upper().strip()

    # Try to extract number from patterns like V12, P3, VISIT 5, etc.
    patterns = [
        r'V(\d+)',      # V12, V4
        r'P(\d+)',      # P3, P11
        r'VISIT\s*(\d+)',  # Visit 12
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return int(match.group(1))

    return None

def normalize_event_code(code):
    """
    Normalize event code for comparison. Returns tuple (normalized_string, visit_number)
    """
    if not code or pd.isna(code):
        return ('', None)

    code = str(code).upper().strip()
    visit_num = extract_visit_number(code)

    # Remove common prefixes and spaces
    normalized = re.sub(r'(VISIT|EVENT)\s*', '', code)

    return (normalized, visit_num)

def check_event_match(event_code, event_name):
    """
    Check if an event code matches an event name using multiple strategies.
    Returns True if they match.
    """
    if not event_code or not event_name:
        return False

    event_code = str(event_code).upper().strip()
    event_name = str(event_name).upper().strip()

    # Strategy 1: Exact match
    if event_code == event_name:
        return True

    # Strategy 2: Substring match
    if event_code in event_name or event_name in event_code:
        return True

    # Strategy 3: Special keywords match (screening, randomisation, etc.)
    special_keywords = {
        'SCREEN': ['SCREEN', 'SCRN', 'SCR'],
        'RANDOM': ['RANDOM', 'RAND', 'RANDOMISATION', 'RANDOMIZATION'],
        'BASELINE': ['BASELINE', 'BASE', 'BL'],
        'FOLLOW': ['FOLLOW', 'FOLLOWUP', 'FU'],
        'EOT': ['EOT', 'END OF TREATMENT', 'ENDOFTREATMENT'],
        'EOS': ['EOS', 'END OF STUDY', 'ENDOFSTUDY'],
    }

    for key, variants in special_keywords.items():
        if any(v in event_code for v in variants) and any(v in event_name for v in variants):
            return True

    # Strategy 4: Visit number match
    code_normalized, code_num = normalize_event_code(event_code)
    name_normalized, name_num = normalize_event_code(event_name)

    if code_num is not None and name_num is not None:
        if code_num == name_num:
            return True

    return False

def merge_crf_into_schedule_grid(schedule_grid_df, crf_raw_df):
    """
    Merge CRF form labels, form codes, and event codes into schedule grid.

    Parameters:
    schedule_grid_df: DataFrame with schedule grid data
    crf_raw_df: DataFrame with raw CRF extraction (includes form_label, form_code, event_codes)

    Returns:
    merged_df: DataFrame with CRF data appended to schedule grid
    """
    try:
        print("[MERGE] Starting CRF merge into schedule grid")

        if schedule_grid_df is None or schedule_grid_df.empty:
            print("[MERGE] Error: Empty schedule grid")
            return schedule_grid_df

        if crf_raw_df is None or crf_raw_df.empty:
            print("[MERGE] Error: Empty CRF data")
            return schedule_grid_df

        # Get unique forms (form_label + form_code combinations)
        required_cols = ['form_label', 'form_code', 'event_codes']
        missing_cols = [col for col in required_cols if col not in crf_raw_df.columns]
        if missing_cols:
            print(f"[MERGE] Warning: Missing columns {missing_cols}, using available data")

        # Extract unique forms with their event codes
        unique_forms = []
        for form_label in crf_raw_df['form_label'].dropna().unique():
            form_data = crf_raw_df[crf_raw_df['form_label'] == form_label].iloc[0]
            form_code = form_data.get('form_code', '') if 'form_code' in crf_raw_df.columns else ''
            event_codes = form_data.get('event_codes', '') if 'event_codes' in crf_raw_df.columns else ''

            unique_forms.append({
                'form_label': form_label,
                'form_code': form_code,
                'event_codes': event_codes
            })

        print(f"[MERGE] Found {len(unique_forms)} unique forms")

        # Find last filled row in schedule grid
        last_filled_row = -1
        for idx in range(len(schedule_grid_df) - 1, -1, -1):
            if schedule_grid_df.iloc[idx].notna().any():
                last_filled_row = idx
                break

        print(f"[MERGE] Last filled row: {last_filled_row}")

        # Create new rows to add
        new_rows = []

        # Add 3 empty rows
        for _ in range(3):
            new_rows.append({col: '' for col in schedule_grid_df.columns})

        # Add header row
        header_row = {'Form Label': 'Form Label'}
        for col in schedule_grid_df.columns[1:]:
            if col == schedule_grid_df.columns[1]:
                header_row[col] = 'Form Code'
            else:
                header_row[col] = ''
        new_rows.append(header_row)

        # Get schedule grid column names for matching
        schedule_columns = schedule_grid_df.columns.tolist()
        print(f"[MERGE] Schedule grid has {len(schedule_columns)} columns")

        # Get Event Names from row 1 of schedule grid
        event_names_row = {}
        if len(schedule_grid_df) > 1:
            for col in schedule_columns[2:]:  # Skip first two columns (Form Label, Form Code equivalent)
                event_names_row[col] = str(schedule_grid_df.iloc[1][col]).strip()

        # Add form rows with X marks
        for form in unique_forms:
            form_row = {
                'Form Label': form['form_label']
            }

            # Add form code to second column
            form_row[schedule_columns[1]] = form['form_code']

            # Parse event codes and mark X in matching columns
            event_codes_str = str(form['event_codes']) if form['event_codes'] else ''
            event_codes_list = []

            if event_codes_str and event_codes_str != 'nan' and event_codes_str != '':
                # Split by comma, semicolon, or space
                event_codes_list = re.split(r'[,;\s]+', event_codes_str)
                event_codes_list = [code.strip() for code in event_codes_list if code.strip()]

            print(f"[MERGE] Form '{form['form_label']}' has event codes: {event_codes_list}")

            # Strategy: Try to match event codes with event names
            matched_columns = set()

            if event_codes_list:
                for event_code in event_codes_list:
                    for col_name in schedule_columns[2:]:
                        event_name = event_names_row.get(col_name, '')

                        if check_event_match(event_code, event_name):
                            matched_columns.add(col_name)
                            print(f"[MERGE]   Matched '{event_code}' -> '{event_name}' in column '{col_name}'")

            # Fill form row with X marks for matched columns
            for col_name in schedule_columns[2:]:
                if col_name in matched_columns:
                    form_row[col_name] = 'X'
                else:
                    form_row[col_name] = ''

            # If no matches found but we have event codes, log a warning
            if event_codes_list and not matched_columns:
                print(f"[MERGE]   WARNING: No matches found for form '{form['form_label']}' with codes {event_codes_list}")

            new_rows.append(form_row)

        # Create DataFrame from new rows and concatenate
        new_data_df = pd.DataFrame(new_rows)
        merged_df = pd.concat([schedule_grid_df, new_data_df], ignore_index=True)

        print(f"[MERGE] Successfully merged. Total rows: {len(merged_df)}")
        return merged_df

    except Exception as e:
        print(f"[MERGE] Error: {e}")
        import traceback
        traceback.print_exc()
        return schedule_grid_df

# ==================== STREAMLIT UI ====================

# Initialize session state
if 'protocol_df' not in st.session_state:
    st.session_state.protocol_df = None
if 'schedule_grid_df' not in st.session_state:
    st.session_state.schedule_grid_df = None
if 'crf_df' not in st.session_state:
    st.session_state.crf_df = None
if 'crf_raw_df' not in st.session_state:
    st.session_state.crf_raw_df = None
if 'combined_excel' not in st.session_state:
    st.session_state.combined_excel = None
if 'protocol_ready' not in st.session_state:
    st.session_state.protocol_ready = False
if 'schedule_grid_ready' not in st.session_state:
    st.session_state.schedule_grid_ready = False
if 'crf_ready' not in st.session_state:
    st.session_state.crf_ready = False
if 'protocol_error' not in st.session_state:
    st.session_state.protocol_error = None
if 'crf_error' not in st.session_state:
    st.session_state.crf_error = None

# Header
st.title("🏥 Clinical Document Processor")
st.markdown("Upload your **Mock CRF** (.docx) and **Protocol REF** (.pdf) documents to extract and process data.")

# File uploaders
col1, col2 = st.columns(2)
with col1:
    uploaded_crf_file = st.file_uploader("📄 Upload Mock CRF (.docx)", type="docx")
with col2:
    uploaded_protocol_file = st.file_uploader("📋 Upload Protocol REF (.pdf)", type="pdf")

# Process button
if st.button("🚀 Process Documents", type="primary", use_container_width=True):
    if uploaded_crf_file and uploaded_protocol_file:
        
        # Initialize metrics
        metrics = ProcessingMetrics()
        metrics.start_processing()
        
        # Save uploaded files
        crf_path = "temp_crf.docx"
        protocol_path = "temp_protocol.pdf"
        
        with open(crf_path, "wb") as f:
            f.write(uploaded_crf_file.getbuffer())
        with open(protocol_path, "wb") as f:
            f.write(uploaded_protocol_file.getbuffer())
        
        st.success("✅ Files uploaded successfully")
        print(f"[MAIN] Files uploaded: {uploaded_crf_file.name}, {uploaded_protocol_file.name}")
        
        # Reset states
        st.session_state.protocol_ready = False
        st.session_state.schedule_grid_ready = False
        st.session_state.crf_ready = False
        st.session_state.protocol_error = None
        st.session_state.crf_error = None
        st.session_state.crf_raw_df = None
        st.session_state.combined_excel = None
        
        # Create dashboard placeholder
        dashboard_placeholder = st.empty()
        with dashboard_placeholder.container():
            create_dashboard(metrics)
        
        # Process Protocol
        st.markdown("---")
        st.subheader("📊 Protocol REF Processing")
        
        try:
            metrics.protocol_start = time.time()
            print("[MAIN] Starting Protocol processing")
            
            with st.spinner("🔍 Identifying Schedule of Activities tables..."):
                extracted_pdf_path = extract_table_pages(protocol_path)
                st.write('Extracted table pages')
                print("Extracted table pages")
            
            if extracted_pdf_path:
                protocol_df = process_protocol_pdf_pdfplumber(
                    extracted_pdf_path,
                    system_prompt_pr,
                    metrics,
                    dashboard_placeholder
                )
                # st.write(protocol_df)
                metrics.protocol_end = time.time()
                
                if not protocol_df.empty:
                    st.success(f"✅ Protocol: Extracted {len(protocol_df)} rows in {format_time(metrics.get_protocol_time())}")
                    st.session_state.protocol_df = protocol_df
                    st.session_state.protocol_ready = True

                    # Post-process to create schedule grid
                    with st.spinner("🔄 Creating Schedule Grid..."):
                        try:
                            print("[MAIN] Starting schedule grid post-processing")
                            schedule_grid_df = map_protocol_to_schedule_grid(protocol_df)

                            if not schedule_grid_df.empty:
                                st.session_state.schedule_grid_df = schedule_grid_df
                                st.session_state.schedule_grid_ready = True
                                st.success(f"✅ Schedule Grid: Created with {len(schedule_grid_df)} rows")
                                print(f"[MAIN] Schedule grid created successfully with {len(schedule_grid_df)} rows")
                            else:
                                st.warning("⚠️ Schedule Grid could not be created - check if protocol data has required format")
                                print("[MAIN] Schedule grid creation returned empty DataFrame")
                        except Exception as e:
                            st.warning(f"⚠️ Schedule Grid creation failed: {e}")
                            print(f"[MAIN] Schedule grid error: {e}")

                    with st.expander("👁️ Preview Protocol Data", expanded=False):
                        try:
                            st.dataframe(protocol_df.head(20), use_container_width=True)
                        except Exception as e:
                            dup1 = protocol_df.copy()
                            dup1.columns = [f"{c}_{i}" for i,c in enumerate(dup1.columns)]
                            st.dataframe(dup1.head(20), use_container_width=True)

                    if st.session_state.schedule_grid_ready:
                        with st.expander("👁️ Preview Schedule Grid", expanded=False):
                            st.dataframe(st.session_state.schedule_grid_df.head(20), use_container_width=True)

                    # Cleanup
                    if os.path.exists(extracted_pdf_path):
                        os.remove(extracted_pdf_path)
                else:
                    st.warning("⚠️ No Protocol data extracted")
            else:
                st.error("❌ Could not identify Schedule of Activities pages")
                
        except Exception as e:
            metrics.protocol_end = time.time()
            st.session_state.protocol_error = str(e)
            st.error(f"❌ Protocol processing failed: {e}")
            print(f"[MAIN] Protocol error: {e}")
        
        # Update dashboard
        dashboard_placeholder.empty()
        with dashboard_placeholder.container():
            create_dashboard(metrics)
        
        # Process Mock CRF
        st.markdown("---")
        st.subheader("📝 Mock CRF Processing")
        
        try:
            metrics.crf_start = time.time()
            print("[MAIN] Starting CRF processing")
            
            with st.spinner("🧩 Chunking document..."):
                crf_chunks = process_crf_docx(crf_path)
                metrics.chunks_created = len(crf_chunks)
                print(f"[MAIN] Created {len(crf_chunks)} chunks")
            
            st.info(f"📦 Created {len(crf_chunks)} chunks")
            
            # Update dashboard
            dashboard_placeholder.empty()
            with dashboard_placeholder.container():
                create_dashboard(metrics)
            
            with st.spinner("🤖 Extracting CRF data with AI..."):
                crf_df = ai_extract(crf_chunks, System_prompt, metrics, dashboard_placeholder)
            
            metrics.crf_end = time.time()
            
            if not crf_df.empty:
                st.success(f"✅ CRF: Extracted {crf_df.shape[0]} items in {format_time(metrics.get_crf_time())}")

                # Store raw CRF data (includes form_code and event_codes)
                st.session_state.crf_raw_df = crf_df.copy()

                # Map to template (excludes form_code and event_codes)
                final_crf = map_data_manually(crf_df, header_row=4, start_row=5)
                st.session_state.crf_df = final_crf
                st.session_state.crf_ready = True

                # If schedule grid is ready, merge CRF data into it
                if st.session_state.schedule_grid_ready:
                    with st.spinner("🔗 Merging CRF data into Schedule Grid..."):
                        try:
                            print("[MAIN] Starting CRF to Schedule Grid merge")
                            merged_schedule = merge_crf_into_schedule_grid(
                                st.session_state.schedule_grid_df,
                                st.session_state.crf_raw_df
                            )
                            if not merged_schedule.empty:
                                st.session_state.schedule_grid_df = merged_schedule
                                st.success("✅ CRF data merged into Schedule Grid")
                                print("[MAIN] Merge complete")
                            else:
                                st.warning("⚠️ Could not merge CRF data into Schedule Grid")
                        except Exception as e:
                            st.warning(f"⚠️ Merge failed: {e}")
                            print(f"[MAIN] Merge error: {e}")

                    # Create combined Excel file with both sheets
                    with st.spinner("📊 Creating combined Excel file..."):
                        try:
                            print("[MAIN] Creating combined Excel file")
                            excel_data = create_combined_excel(
                                st.session_state.schedule_grid_df,
                                st.session_state.crf_df
                            )
                            if excel_data:
                                st.session_state.combined_excel = excel_data
                                st.success("✅ Combined Excel file created")
                                print("[MAIN] Combined Excel created successfully")
                            else:
                                st.warning("⚠️ Could not create combined Excel file")
                        except Exception as e:
                            st.warning(f"⚠️ Excel creation failed: {e}")
                            print(f"[MAIN] Excel creation error: {e}")

                with st.expander("👁️ Preview CRF Data", expanded=False):
                    st.markdown("**Raw Extracted Data:**")
                    st.dataframe(crf_df.head(20), use_container_width=True)
                    st.markdown("**Mapped to Template:**")
                    st.dataframe(final_crf.head(20), use_container_width=True)
            else:
                st.warning("⚠️ No CRF data extracted")
                
        except Exception as e:
            metrics.crf_end = time.time()
            st.session_state.crf_error = str(e)
            st.error(f"❌ CRF processing failed: {e}")
            print(f"[MAIN] CRF error: {e}")
        
        # Final dashboard update
        dashboard_placeholder.empty()
        with dashboard_placeholder.container():
            create_dashboard(metrics)
        
        # Cleanup
        for f in [crf_path, protocol_path]:
            if os.path.exists(f):
                os.remove(f)
        
        # Summary
        st.markdown("---")
        if st.session_state.protocol_ready or st.session_state.schedule_grid_ready or st.session_state.crf_ready:
            st.success("🎉 Processing complete! Download your results below.")
        else:
            st.error("❌ Both processes failed. Please check your files and try again.")

# ==================== DOWNLOAD SECTIONS ====================
st.markdown("---")
st.header("📥 Download Results")

# Download 1 - Combined Excel File (PRIMARY - RECOMMENDED)
if st.session_state.combined_excel is not None:
    st.markdown("### 1️⃣ Combined Excel File ⭐ (Recommended)")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.success("✅ Complete dataset with both Schedule Grid and CRF data")
        st.info("📊 Contains 2 sheets: 'Schedule Grid' (with merged CRF forms) + 'Study Specific Forms'")
    with col2:
        st.download_button(
            "📥 Download Combined Excel",
            data=st.session_state.combined_excel,
            file_name=f'clinical_study_data_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            type='primary',
            use_container_width=True,
            key='download_combined_excel'
        )
    st.markdown("---")

# Download 2 - Protocol (AI Processed Tables)
if st.session_state.protocol_ready:
    st.markdown("### 2️⃣ Protocol REF - AI Processed Tables (CSV)")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.success(f"✅ Protocol extracted data ready ({len(st.session_state.protocol_df)} rows)")
        st.info("📄 Raw extracted Schedule of Activities tables from Protocol REF")
    with col2:
        st.download_button(
            "📥 Download Protocol CSV",
            data=st.session_state.protocol_df.to_csv(index=False).encode('utf-8'),
            file_name=f'protocol_extraction_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv',
            mime='text/csv',
            use_container_width=True,
            key='download_protocol_csv'
        )
    st.markdown("---")
elif st.session_state.protocol_error:
    st.markdown("### 2️⃣ Protocol REF - AI Processed Tables")
    st.error(f"❌ Protocol error: {st.session_state.protocol_error}")
    st.markdown("---")

# Download 3 - Schedule Grid (Mapped Protocol)
if st.session_state.schedule_grid_ready:
    st.markdown("### 3️⃣ Schedule Grid - Mapped Protocol (CSV)")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.success(f"✅ Schedule Grid mapping ready ({len(st.session_state.schedule_grid_df)} rows)")
        st.info("📄 Protocol data mapped to Schedule Grid format with CRF forms merged")
    with col2:
        st.download_button(
            "📥 Download Schedule Grid CSV",
            data=st.session_state.schedule_grid_df.to_csv(index=False).encode('utf-8'),
            file_name=f'schedule_grid_mapped_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv',
            mime='text/csv',
            use_container_width=True,
            key='download_schedule_csv'
        )
    st.markdown("---")

# Download 4 - CRF Output
if st.session_state.crf_ready:
    st.markdown("### 4️⃣ CRF - Extracted Data (CSV)")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.success(f"✅ CRF extraction ready ({len(st.session_state.crf_df)} rows)")
        st.info("📄 Extracted and mapped CRF data from Mock CRF document")
    with col2:
        st.download_button(
            "📥 Download CRF CSV",
            data=st.session_state.crf_df.to_csv(index=False).encode('utf-8'),
            file_name=f'crf_extraction_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv',
            mime='text/csv',
            use_container_width=True,
            key='download_crf_csv'
        )
    st.markdown("---")
elif st.session_state.crf_error:
    st.markdown("### 4️⃣ CRF - Extracted Data")
    st.error(f"❌ CRF error: {st.session_state.crf_error}")
    st.markdown("---")

# Reset button
if (st.session_state.protocol_ready or st.session_state.schedule_grid_ready or st.session_state.crf_ready or
    st.session_state.protocol_error or st.session_state.crf_error):
    st.markdown("---")
    if st.button('🔄 Process New Documents', use_container_width=True):
        # Clear session state
        st.session_state.protocol_df = None
        st.session_state.schedule_grid_df = None
        st.session_state.crf_df = None
        st.session_state.crf_raw_df = None
        st.session_state.combined_excel = None
        st.session_state.protocol_ready = False
        st.session_state.schedule_grid_ready = False
        st.session_state.crf_ready = False
        st.session_state.protocol_error = None
        st.session_state.crf_error = None

        # Clean up temp files
        for f in ["temp_crf.docx", "temp_protocol.pdf", "Schedule_of_Activities.pdf"]:
            if os.path.exists(f):
                os.remove(f)

        print("[MAIN] Reset complete, ready for new documents")
        st.rerun()

# Footer
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: #666; padding: 20px;'>"
    "Clinical Document Processor | Built with Streamlit & OpenAI"
    "</div>",
    unsafe_allow_html=True
)
