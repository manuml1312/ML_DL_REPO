import streamlit as st
import re
import fitz
import pandas as pd
from docx import Document
from typing import List, Dict, Any
from openai import OpenAI
import json
import os
import pdfplumber
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

# API setup
api_key = st.secrets["api_key"]
client = OpenAI(api_key=api_key)

# Thread-safe progress tracking
class ProgressTracker:
    def __init__(self, total_items, description):
        self.lock = threading.Lock()
        self.completed = 0
        self.total = total_items
        self.progress_bar = st.progress(0, text=f"{description} (0/{total_items})")
    
    def update(self, increment=1):
        with self.lock:
            self.completed += increment
            progress = self.completed / self.total
            self.progress_bar.progress(
                progress, 
                text=f"Processing ({self.completed}/{self.total})"
            )
    
    def complete(self):
        self.progress_bar.empty()


class DOCXCRFChunker:
    def __init__(self, max_chunk_size: int = 15000, overlap_size: int = 500):
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size
        self.form_patterns = [
            r'sCRF\s+v\d+\.\d+',
            r'\[[\w_]+\]',
            r'V\d+(?:,\s*V\d+)*',
        ]

    def extract_and_chunk(self, docx_path: str) -> List[Dict[str, Any]]:
        doc = Document(docx_path)
        elements = self._extract_structured_content(doc)
        chunks = self._create_chunks(elements)
        return chunks

    def _extract_structured_content(self, doc: Document) -> List[Dict[str, Any]]:
        elements = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = self._get_paragraph_from_element(doc, element)
                if para and para.text.strip():
                    elements.append({
                        'type': 'paragraph',
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': self._is_heading(para),
                        'is_form_title': self._is_form_title(para.text),
                        'length': len(para.text)
                    })
            elif element.tag.endswith('tbl'):
                table = self._get_table_from_element(doc, element)
                if table:
                    table_text = self._extract_table_text(table)
                    elements.append({
                        'type': 'table',
                        'text': table_text,
                        'rows': len(table.rows),
                        'cols': len(table.columns) if table.rows else 0,
                        'length': len(table_text),
                        'is_crf_table': self._is_crf_table(table_text)
                    })
        return elements

    def _get_paragraph_from_element(self, doc, element):
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _get_table_from_element(self, doc, element):
        for table in doc.tables:
            if table._element == element:
                return table
        return None

    def _is_heading(self, para) -> bool:
        if para.style:
            style_name = para.style.name.lower()
            return 'heading' in style_name or 'title' in style_name
        return False

    def _is_form_title(self, text: str) -> bool:
        for pattern in self.form_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        form_indicators = [
            'collection of samples', 'contraception', 'c-ssrs',
            'patient health questionnaire', 'weight history',
            'hand grip test', 'mri scan consent'
        ]
        text_lower = text.lower()
        return any(indicator in text_lower for indicator in form_indicators)

    def _extract_table_text(self, table) -> str:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))
        return "\n".join(table_text)

    def _is_crf_table(self, table_text: str) -> bool:
        crf_indicators = [
            'yes', 'no', 'radio', 'checkbox', 'required', 'optional', '*',
            'integration', 'argus', 'cosmos', 'item label', 'data type', 'codelist'
        ]
        text_lower = table_text.lower()
        return sum(1 for indicator in crf_indicators if indicator in text_lower) >= 2

    def _create_chunks(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        chunks = []
        current_chunk = self._create_new_chunk(1)

        for element in elements:
            if (element['type'] == 'paragraph' and element.get('is_form_title', False)):
                if current_chunk['length'] > 0:
                    chunks.append(self._finalize_chunk(current_chunk))
                    current_chunk = self._create_new_chunk(len(chunks) + 1)
                current_chunk['form_context'] = element['text']

            if (current_chunk['length'] + element['length'] > self.max_chunk_size and 
                current_chunk['length'] > 0):
                if element['type'] == 'table':
                    chunks.append(self._finalize_chunk(current_chunk))
                    current_chunk = self._create_new_chunk(len(chunks) + 1)
                    current_chunk['form_context'] = chunks[-1]['form_context'] if chunks else ''

            current_chunk['elements'].append(element)
            current_chunk['text'] += element['text'] + '\n\n'
            current_chunk['length'] += element['length']
            if element['type'] == 'table':
                current_chunk['has_tables'] = True

        if current_chunk['length'] > 0:
            chunks.append(self._finalize_chunk(current_chunk))

        return self._add_overlap(chunks)

    def _create_new_chunk(self, chunk_id: int) -> Dict[str, Any]:
        return {
            'chunk_id': chunk_id,
            'elements': [],
            'text': '',
            'length': 0,
            'form_context': '',
            'has_tables': False
        }

    def _finalize_chunk(self, chunk: Dict[str, Any]) -> Dict[str, Any]:
        if chunk['form_context'] and chunk['form_context'] not in chunk['text'][:200]:
            chunk['text'] = f"FORM CONTEXT: {chunk['form_context']}\n\n" + chunk['text']
        
        chunk['paragraph_count'] = sum(1 for el in chunk['elements'] if el['type'] == 'paragraph')
        chunk['table_count'] = sum(1 for el in chunk['elements'] if el['type'] == 'table')
        chunk['crf_table_count'] = sum(1 for el in chunk['elements'] 
                                      if el['type'] == 'table' and el.get('is_crf_table', False))
        return chunk

    def _add_overlap(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if len(chunks) <= 1:
            return chunks
        
        for i in range(1, len(chunks)):
            prev_text = chunks[i-1]['text']
            overlap_text = prev_text[-self.overlap_size:] if len(prev_text) > self.overlap_size else prev_text
            chunks[i]['text'] = f"OVERLAP FROM PREVIOUS:\n{overlap_text}\n\nCURRENT CHUNK:\n" + chunks[i]['text']
            chunks[i]['has_overlap'] = True
        
        return chunks


# Optimized Prompts
SYSTEM_PROMPT_CRF = """You are a CRF extraction specialist. Extract structured information from clinical research documents and return ONLY valid JSON.

OUTPUT FORMAT:
{
  "forms": [{
    "form_label": "string",
    "form_code": "string",
    "visits": "string",
    "form_type": "string",
    "item_group": "string",
    "item_group_repeating": "Y/N",
    "item_order": number,
    "item_label": "string",
    "item_name": "[SDTM Programmer]",
    "data_type": "string",
    "codelist": "string",
    "codelist_name": "string",
    "control_time": "",
    "required": "Y/N"
  }],
  "chunk_metadata": {
    "chunk_id": "string",
    "forms_found": number,
    "items_extracted": number
  }
}

EXTRACTION RULES:
1. Form Label: Main form title (e.g., "Collection of Samples for Laboratory (Lab_1)")
2. Form Code: Code in brackets (e.g., "[LAB_SMPL_TKN]")
3. Visits: Visit schedule (e.g., "V1, V4, V5, V6")
4. Form Type: "Non-repeating form" or "Repeating form"
5. Item Group: Sections like "Blood", "Urine", "Contraception"
6. Item Group Repeating: "Y" if repeating, "N" otherwise
7. Item Order: Sequential number (1, 2, 3...)
8. Item Label: Exact question text
9. Item Name: Always "[SDTM Programmer]"
10. Data Type: Radio Button/Numeric/Text/Date/Checkbox
11. Codelist: Available options (e.g., "Yes/No")
12. Codelist Name: Logical name for options
13. Required: "Y" if asterisk (*) present, "N" otherwise

Return ONLY valid JSON, no explanations."""

SYSTEM_PROMPT_PROTOCOL = """You are a clinical trial data structuring specialist. Clean and restructure Schedule of Activities table JSON.

INPUT: Messy JSON with multiple visit codes packed into single cells, merged phase names, misaligned timing data.

REQUIRED TRANSFORMATIONS:
1. Split Merged Visit Columns - Create separate columns for each visit (e.g., "V2D-2\nV2D-1 V2D1" → 3 columns)
2. Propagate Phase Names - Fill null values with phase name from row above
3. Clean Text - Remove "\n", fix spacing, keep protocol section references
4. Align Timing Data - Ensure days/weeks/windows align with correct visits
5. Preserve Structure - Maintain row order, keep headers, preserve all "X" marks

OUTPUT: Clean JSON with each visit in own column, phase names filled, text cleaned, timing aligned, X marks preserved.

CRITICAL: When splitting visits, X marks and timing values must stay with correct visit.

Return ONLY cleaned JSON, no explanations."""


def process_crf_chunk(chunk, chunk_idx, total_chunks):
    """Process a single CRF chunk"""
    try:
        user_prompt = f"""Extract CRF information from this document chunk and return JSON:

CHUNK CONTENT:
{chunk['text']}

Analyze and extract all CRF forms, item groups, and items following the specified JSON format. Focus on:
- Form titles and codes
- Required fields (marked with *)
- Item groups and organization
- Question text and response options
- Data types

Return only valid JSON."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT_CRF},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        
        answer = json.loads(response.choices[0].message.content)
        return answer.get('forms', [])
    
    except Exception as e:
        st.error(f"Error processing CRF chunk {chunk_idx + 1}/{total_chunks}: {e}")
        return []


def process_protocol_table(page_idx, page, table_idx, table_data, total_pages):
    """Process a single protocol table"""
    try:
        raw_data = [list(row) for row in table_data]
        raw_json = json.dumps({"data": raw_data})
        
        user_prompt = f"""INPUT JSON: {raw_json}

Clean and return the structured JSON."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT_PROTOCOL},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        
        cleaned_data = json.loads(response.choices[0].message.content)
        return cleaned_data.get('data', [])
    
    except Exception as e:
        st.error(f"Error processing protocol table on page {page_idx + 1}: {e}")
        return []


def process_crf_parallel(chunks):
    """Process CRF chunks in parallel"""
    all_forms_data = []
    tracker = ProgressTracker(len(chunks), "Processing CRF chunks")
    
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {
            executor.submit(process_crf_chunk, chunk, i, len(chunks)): i 
            for i, chunk in enumerate(chunks)
        }
        
        for future in as_completed(futures):
            forms = future.result()
            if forms:
                all_forms_data.extend(forms)
            tracker.update()
    
    tracker.complete()
    return pd.DataFrame(all_forms_data) if all_forms_data else pd.DataFrame()


def process_protocol_parallel(pdf_path):
    """Process protocol PDF tables in parallel"""
    table_settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "snap_tolerance": 300,
        "snap_x_tolerance": 6,
        "snap_y_tolerance": 5.16,
        "join_tolerance": 1,
        "join_x_tolerance": 2,
        "join_y_tolerance": 3,
        "edge_min_length": 100,
        "min_words_vertical": 3,
        "min_words_horizontal": 1,
        "intersection_tolerance": 1,
        "intersection_x_tolerance": 0.4,
        "intersection_y_tolerance": 2,
        "text_tolerance": 3,
        "text_x_tolerance": 5,
        "text_y_tolerance": 3,
    }
    
    all_extracted_data = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # Collect all tables first
            tables_to_process = []
            for page_idx, page in enumerate(pdf.pages):
                tables_on_page = page.extract_tables(table_settings=table_settings)
                if tables_on_page:
                    for table_idx, table_data in enumerate(tables_on_page):
                        tables_to_process.append((page_idx, page, table_idx, table_data, len(pdf.pages)))
            
            if not tables_to_process:
                st.warning("No tables found in Protocol PDF")
                return pd.DataFrame()
            
            tracker = ProgressTracker(len(tables_to_process), "Processing Protocol tables")
            
            # Process tables in parallel
            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = {
                    executor.submit(process_protocol_table, *args): i 
                    for i, args in enumerate(tables_to_process)
                }
                
                for future in as_completed(futures):
                    cleaned_data = future.result()
                    if cleaned_data:
                        all_extracted_data.extend(cleaned_data)
                    tracker.update()
            
            tracker.complete()
            
            if all_extracted_data:
                pr_df = pd.DataFrame(all_extracted_data)
                
                if not pr_df.empty:
                    pr_df.columns = pr_df.iloc[0]
                    pr_df = pr_df[1:].reset_index(drop=True)
                
                new_cols = ['Common Forms', 'Unscheduled', 'Is Form Dynamic?', 
                           'Form Dynamic Criteria', 'Additional Programming Instructions']
                for col in new_cols:
                    if col not in pr_df.columns:
                        pr_df[col] = ''
                
                pr_df = pr_df.dropna(axis=1, how='all')
                return pr_df
            
            return pd.DataFrame()
    
    except Exception as e:
        st.error(f"Error processing Protocol PDF: {e}")
        return pd.DataFrame()


# Streamlit App
st.set_page_config(page_title="Clinical Document Processor", layout="wide")
st.title("📋 Clinical Document Processor")
st.write("Upload your Mock CRF (.docx) and Protocol REF (.pdf) to extract data in parallel.")

col1, col2 = st.columns(2)

with col1:
    uploaded_crf_file = st.file_uploader("Upload Mock CRF (.docx)", type="docx")

with col2:
    uploaded_protocol_file = st.file_uploader("Upload Protocol REF (.pdf)", type="pdf")

if st.button("🚀 Process Documents", type="primary"):
    if uploaded_crf_file and uploaded_protocol_file:
        
        # Save files temporarily
        crf_path = "temp_crf.docx"
        protocol_path = "temp_protocol.pdf"
        
        with open(crf_path, "wb") as f:
            f.write(uploaded_crf_file.getbuffer())
        with open(protocol_path, "wb") as f:
            f.write(uploaded_protocol_file.getbuffer())
        
        st.success("✅ Files uploaded successfully")
        
        # Create two columns for parallel display
        col_crf, col_protocol = st.columns(2)
        
        with col_crf:
            st.subheader("📄 Mock CRF Processing")
            with st.spinner("Chunking CRF document..."):
                chunker = DOCXCRFChunker(max_chunk_size=15000, overlap_size=500)
                crf_chunks = chunker.extract_and_chunk(crf_path)
                st.info(f"Created {len(crf_chunks)} chunks")
        
        # Process both in parallel using separate threads
        crf_result = [None]
        protocol_result = [None]
        
        def process_crf_wrapper():
            crf_result[0] = process_crf_parallel(crf_chunks)
        
        def process_protocol_wrapper():
            protocol_result[0] = process_protocol_parallel(protocol_path)
        
        # Start both processes
        thread_crf = threading.Thread(target=process_crf_wrapper)
        thread_protocol = threading.Thread(target=process_protocol_wrapper)
        
        thread_crf.start()
        thread_protocol.start()
        
        # Wait for both to complete
        thread_crf.join()
        thread_protocol.join()
        
        # Display results
        with col_crf:
            if crf_result[0] is not None and not crf_result[0].empty:
                st.success(f"✅ Extracted {len(crf_result[0])} items")
                st.dataframe(crf_result[0], use_container_width=True)
                
                st.download_button(
                    "📥 Download CRF Data",
                    data=crf_result[0].to_csv(index=False).encode('utf-8'),
                    file_name='crf_extraction.csv',
                    mime='text/csv'
                )
            else:
                st.warning("No CRF data extracted")
        
        with col_protocol:
            st.subheader("📊 Protocol REF Processing")
            if protocol_result[0] is not None and not protocol_result[0].empty:
                st.success(f"✅ Extracted {len(protocol_result[0])} rows")
                st.dataframe(protocol_result[0], use_container_width=True)
                
                st.download_button(
                    "📥 Download Protocol Data",
                    data=protocol_result[0].to_csv(index=False).encode('utf-8'),
                    file_name='protocol_extraction.csv',
                    mime='text/csv'
                )
            else:
                st.warning("No Protocol data extracted")
        
        # Cleanup
        for f in [crf_path, protocol_path]:
            if os.path.exists(f):
                os.remove(f)
        
        st.success("✅ Processing complete!")
    
    else:
        st.error("⚠️ Please upload both documents")
