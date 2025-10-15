import re
from typing import List, Dict, Any
from docx import Document
import pandas as pd
import json
import streamlit as st
from openai import OpenAI

# Initialize OpenAI client (should be done with your API key)
# client = OpenAI(api_key="your-api-key-here")

class DOCXCRFChunker:
    def __init__(self, max_chunk_size: int = 15000, overlap_size: int = 500, openai_client=None):
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size
        self.client = openai_client  # Store OpenAI client

        # CRF form detection patterns
        self.form_patterns = [
            r'sCRF\s+v\d+\.\d+',  # sCRF v1.0
            r'\[[\w_]+\]',        # [FORM_CODE]
            r'V\d+(?:,\s*V\d+)*', # V1, V2, V3 (visit schedules)
        ]
        
        # System prompt for AI extraction
        self.system_prompt = """You are a CRF (Case Report Form) extraction specialist. Your task is to extract structured information from clinical research document chunks and return it as valid JSON.
REQUIRED OUTPUT FORMAT: Return ONLY a valid JSON object with the structure below. Do not include any explanatory text before or after the JSON.
{
  "forms": [
    {
      "form_label": "string",
      "form_code": "string",
      "visits": "string",
      "form_type": "string",
      "item_group": "string",
      "item_group_repeating": "Y/N",
      "item_order": number,
      "item_label": "string",
      "item_name": "[SDTM Programmer]",
      "data_type": "string",
      "codelist": "string",
      "codelist_name": "string",
      "control_time": "",
      "required": "Y/N"
    }
  ],
  "chunk_metadata": {
    "chunk_id": "string",
    "forms_found": number,
    "items_extracted": number
  }
}

EXTRACTION RULES:
1. Form Label: Extract the main form title (e.g., "Collection of Samples for Laboratory (Lab_1)")
2. Form Code: Extract code in brackets (e.g., "[LAB_SMPL_TKN]")
3. Visits: Extract visit schedule (e.g., "V1, V4, V5, V6")
4. Form Type: "Non-repeating form" or "Repeating form"
5. Item Group: Logical sections like "Blood", "Urine", "Contraception"
6. Item Group Repeating: "Y" if section can repeat, "N" otherwise
7. Item Order: Sequential number within form (1, 2, 3...)
8. Item Label: Exact question text
9. Item Name: Always "[SDTM Programmer]"
10. Data Type: "Radio Button", "Numeric", "Text", "Date", "Checkbox"
11. Codelist: Available options (e.g., "Yes/No", "4-point Scale")
12. Codelist Name: Logical name for options
13. Required: "Y" if marked with asterisk (*), "N" otherwise

DATA TYPE MAPPING:
- Questions with radio options → "Radio Button"
- Numeric inputs (age, weight, scores) → "Numeric"
- Free text fields → "Text"
- Date fields → "Date"
- Multiple selections → "Checkbox"

Return valid JSON only. No other text."""

    def extract_and_chunk(self, docx_path: str):
        """Extract text from DOCX and create table-aware chunks"""
        doc = Document(docx_path)
        elements = self._extract_structured_content(doc)
        chunks = self._create_chunks(elements)
        return chunks

    def _extract_structured_content(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract content while preserving structure"""
        elements = []

        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = self._get_paragraph_from_element(doc, element)
                if para and para.text.strip():
                    elements.append({
                        'type': 'paragraph',
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': self._is_heading(para),
                        'is_form_title': self._is_form_title(para.text),
                        'length': len(para.text)
                    })

            elif element.tag.endswith('tbl'):  # Table
                table = self._get_table_from_element(doc, element)
                if table:
                    table_text = self._extract_table_text(table)
                    elements.append({
                        'type': 'table',
                        'text': table_text,
                        'rows': len(table.rows),
                        'cols': len(table.columns) if table.rows else 0,
                        'length': len(table_text),
                        'is_crf_table': self._is_crf_table(table_text)
                    })

        return elements

    def _get_paragraph_from_element(self, doc, element):
        """Get paragraph object from XML element"""
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _get_table_from_element(self, doc, element):
        """Get table object from XML element"""
        for table in doc.tables:
            if table._element == element:
                return table
        return None

    def _is_heading(self, para) -> bool:
        """Check if paragraph is a heading"""
        if para.style:
            style_name = para.style.name.lower()
            return 'heading' in style_name or 'title' in style_name
        return False

    def _is_form_title(self, text: str) -> bool:
        """Check if text is a CRF form title"""
        for pattern in self.form_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True

        # Additional form title indicators
        form_indicators = [
            'collection of samples',
            'contraception',
            'c-ssrs',
            'patient health questionnaire',
            'weight history',
            'hand grip test',
            'mri scan consent'
        ]

        text_lower = text.lower()
        return any(indicator in text_lower for indicator in form_indicators)

    def _extract_table_text(self, table) -> str:
        """Extract text from table preserving structure"""
        table_text = []

        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                table_text.append(" | ".join(row_text))

        return "\n".join(table_text)

    def _is_crf_table(self, table_text: str) -> bool:
        """Check if table contains CRF content"""
        crf_indicators = [
            'yes', 'no', 'radio', 'checkbox',
            'required', 'optional', '*',
            'integration', 'argus', 'cosmos',
            'item label', 'data type', 'codelist'
        ]

        text_lower = table_text.lower()
        return sum(1 for indicator in crf_indicators if indicator in text_lower) >= 2

    def _create_chunks(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Create table-aware chunks from extracted elements"""
        chunks = []
        current_chunk = {
            'chunk_id': 1,
            'elements': [],
            'text': '',
            'length': 0,
            'form_context': '',
            'has_tables': False
        }

        for element in elements:
            # Check if this element starts a new form
            if (element['type'] == 'paragraph' and
                element.get('is_form_title', False)):

                # Finish current chunk if it has content
                if current_chunk['length'] > 0:
                    chunks.append(self._finalize_chunk(current_chunk))
                    current_chunk = self._create_new_chunk(len(chunks) + 1)

                # Set form context for new chunk
                current_chunk['form_context'] = element['text']

            # Check if adding this element exceeds chunk size
            if (current_chunk['length'] + element['length'] > self.max_chunk_size and
                current_chunk['length'] > 0):

                # Don't break in the middle of a table
                if element['type'] == 'table':
                    # Finish current chunk and start new one
                    chunks.append(self._finalize_chunk(current_chunk))
                    current_chunk = self._create_new_chunk(len(chunks) + 1)
                    current_chunk['form_context'] = chunks[-1]['form_context']

            # Add element to current chunk
            current_chunk['elements'].append(element)
            current_chunk['text'] += element['text'] + '\n\n'
            current_chunk['length'] += element['length']

            if element['type'] == 'table':
                current_chunk['has_tables'] = True

        # Add final chunk
        if current_chunk['length'] > 0:
            chunks.append(self._finalize_chunk(current_chunk))

        # Add overlap between chunks
        chunks = self._add_overlap(chunks)

        return chunks

    def _create_new_chunk(self, chunk_id: int) -> Dict[str, Any]:
        """Create a new empty chunk"""
        return {
            'chunk_id': chunk_id,
            'elements': [],
            'text': '',
            'length': 0,
            'form_context': '',
            'has_tables': False
        }

    def _finalize_chunk(self, chunk: Dict[str, Any]) -> Dict[str, Any]:
        """Finalize chunk with metadata"""
        # Add form context at the beginning if not already there
        if chunk['form_context'] and chunk['form_context'] not in chunk['text'][:200]:
            chunk['text'] = f"FORM CONTEXT: {chunk['form_context']}\n\n" + chunk['text']

        # Calculate statistics
        chunk['paragraph_count'] = sum(1 for el in chunk['elements'] if el['type'] == 'paragraph')
        chunk['table_count'] = sum(1 for el in chunk['elements'] if el['type'] == 'table')
        chunk['crf_table_count'] = sum(1 for el in chunk['elements']
                                     if el['type'] == 'table' and el.get('is_crf_table', False))

        return chunk

    def _add_overlap(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Add overlap between consecutive chunks"""
        if len(chunks) <= 1:
            return chunks

        for i in range(1, len(chunks)):
            prev_chunk = chunks[i-1]
            current_chunk = chunks[i]

            # Get overlap text from previous chunk
            prev_text = prev_chunk['text']
            overlap_text = prev_text[-self.overlap_size:] if len(prev_text) > self.overlap_size else prev_text

            # Add overlap to current chunk with marker
            current_chunk['text'] = f"OVERLAP FROM PREVIOUS:\n{overlap_text}\n\nCURRENT CHUNK:\n" + current_chunk['text']
            current_chunk['has_overlap'] = True

        return chunks

    def get_chunk_summary(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get summary statistics for chunks"""
        return {
            'total_chunks': len(chunks),
            'total_length': sum(chunk['length'] for chunk in chunks),
            'avg_chunk_size': sum(chunk['length'] for chunk in chunks) / len(chunks) if chunks else 0,
            'chunks_with_tables': sum(1 for chunk in chunks if chunk['has_tables']),
            'total_tables': sum(chunk.get('table_count', 0) for chunk in chunks),
            'total_crf_tables': sum(chunk.get('crf_table_count', 0) for chunk in chunks),
            'forms_identified': len(set(chunk['form_context'] for chunk in chunks if chunk['form_context']))
        }

    def combine_chunks(self, chunks, max_tokens=1500):
        """
        Combine chunks that fit within max_tokens limit
        
        Args:
            chunks: List of chunk dictionaries
            max_tokens: Maximum token limit for combined chunks
        
        Returns:
            List of combined chunk dictionaries
        """
        if not chunks or len(chunks) == 0:
            return []
        
        updated_chunks = []
        current_chunk = chunks[0].copy()
        
        for i in range(1, len(chunks)):
            try:
                next_chunk = chunks[i]
                
                # Calculate combined length with 25% overhead
                combined_length = current_chunk.get('length', 0) + next_chunk.get('length', 0) * 1.25
                
                if combined_length < max_tokens:
                    # Merge chunks
                    current_chunk = {
                        'chunk_id': f"{current_chunk.get('chunk_id', '')}-{next_chunk.get('chunk_id', '')}",
                        'elements': current_chunk.get('elements', []) + next_chunk.get('elements', []),
                        'text': current_chunk.get('text', '') + "\n\n" + next_chunk.get('text', ''),
                        'length': current_chunk.get('length', 0) + next_chunk.get('length', 0),
                        'form_context': current_chunk.get('form_context', '') + '\n\n' + next_chunk.get('form_context', ''), 
                        'has_tables': current_chunk.get('has_tables', False) or next_chunk.get('has_tables', False),
                        'paragraph_count': current_chunk.get('paragraph_count', 0) + next_chunk.get('paragraph_count', 0),
                        'table_count': current_chunk.get('table_count', 0) + next_chunk.get('table_count', 0),
                        'crf_table_count': current_chunk.get('crf_table_count', 0) + next_chunk.get('crf_table_count', 0),
                        'has_overlap': current_chunk.get('has_overlap', False) or next_chunk.get('has_overlap', False)
                    }
                else:
                    # Save current chunk and start new one
                    updated_chunks.append(current_chunk)
                    current_chunk = next_chunk.copy()
                    
            except Exception as e:
                st.write(f"⚠️ Error processing chunk {i}: {e}")
                # On error, save current and continue with next
                updated_chunks.append(current_chunk)
                current_chunk = chunks[i].copy()
        
        # Don't forget the last chunk!
        updated_chunks.append(current_chunk)
        
        st.write(f"✅ Combined {len(chunks)} chunks into {len(updated_chunks)} chunks")
        return updated_chunks

    def process_crf_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """Process a CRF DOCX file and return chunks"""
        chunks = self.extract_and_chunk(docx_path)
        chunks = self.combine_chunks(chunks, 1500)
    
        # Print summary
        summary = self.get_chunk_summary(chunks)
        st.write(f"Created {summary['total_chunks']} chunks from Mock CRF.")
        st.write(f"Forms identified: {summary['forms_identified']}")
        
        return chunks

    def user_prompt(self, text):
        """Generate user prompt for AI extraction"""
        prompt = f"""Extract CRF information from the following document chunk and return as JSON:

CHUNK CONTENT:
{text}

Analyze this content and extract all CRF forms, item groups, and individual items following the specified JSON format. Pay special attention to:
- Form titles and codes
- Required fields marked with asterisks (*)
- Item groups and their organization
- Question text and response options
- Data types based on field characteristics

Return only valid JSON with no additional text."""
        
        return prompt

    def ai_extract(self, chunks):
        """
        Function to call OpenAI API for each chunk
        
        Args:
            chunks: List of chunk dictionaries
            
        Returns:
            DataFrame containing extracted forms data
        """
        if not self.client:
            st.error("OpenAI client not initialized. Please provide API client.")
            return pd.DataFrame()
            
        conversation = [{"role": "system", "content": self.system_prompt}]
        all_forms_data = []
    
        progress_text = "Operation in progress. Please wait."
        my_bar = st.progress(0, text=progress_text)
    
        for i, chunk in enumerate(chunks):
            user_msg = self.user_prompt(chunk['text'])
            messages = conversation + [{"role": "user", "content": user_msg}]
    
            try:
                response = self.client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=messages,
                    response_format={"type": "json_object"}
                )
                answer = json.loads(response.choices[0].message.content)
    
                if 'forms' in answer:
                    all_forms_data.extend(answer['forms'])
    
                # Update progress bar
                progress_percentage = (i + 1) / len(chunks)
                my_bar.progress(progress_percentage, text=f"Processing chunk {i+1}/{len(chunks)}")
    
            except Exception as e:
                st.error(f"Error processing chunk {chunk['chunk_id']}: {e}")
                st.write("Attempting to continue with next chunk...")
    
        my_bar.empty()  # Clear progress bar on completion
    
        if all_forms_data:
            df = pd.DataFrame(all_forms_data)
            return df
        else:
            st.warning("No forms data extracted from the Mock CRF chunks.")
            return pd.DataFrame()

